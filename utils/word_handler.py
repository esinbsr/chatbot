from docx import Document
from pathlib import Path

# export text to word
def export_to_word(text: str, filename: str = None):
    if filename is None:
       filename = "doc"
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename + ".docx")

# read pdf
def read_word(doc_path:str)-> str:
    doc_path = Path(doc_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {doc_path}")
    try:
        doc = Document( doc_path)
        text = (p.text for p in doc.paragraphs)
        return text
    except Exception as e:
        raise RuntimeError(f"Erreur lors de la lecture du document : {e}")